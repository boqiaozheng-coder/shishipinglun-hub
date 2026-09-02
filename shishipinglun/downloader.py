#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
时评文章下载归档工具

抓取以下来源的最新评论文章，并导出为 PDF / Word 文档：
  1. 人民日报：人民时评、今日谈
  2. 求是网（社论评论）、半月谈网（评论）
  3. The New York Times（Opinion 栏目，RSS）

输出到桌面并按三大类别建档：
  01_人民日报_人民时评与今日谈 /
  02_求是网与半月谈 /
  03_纽约时报观点 /

仅供个人学习、离线阅读使用；请控制抓取频率并尊重各网站版权。
"""

from __future__ import annotations

import argparse
import datetime as dt
import math
import os
import re
import sys
import time
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from email.utils import parsedate_to_datetime
from pathlib import Path
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    HAVE_DOCX = True
except Exception:  # pragma: no cover
    HAVE_DOCX = False

try:
    import pymupdf as fitz
except Exception:  # pragma: no cover
    try:
        import fitz
    except Exception:
        fitz = None


USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/125.0 Safari/537.36"
)

DEFAULT_NYT_RSS = "https://rss.nytimes.com/services/xml/rss/nyt/Opinion.xml"

# 桌面默认输出目录（可在命令行用 --out 覆盖）
DEFAULT_OUT = Path.home() / "Desktop" / "时评文章下载"

last_run_summary = ""  # 供桌面端/服务端读取最近一次下载结果

CATEGORIES = {
    "cat1": "01_人民日报_人民时评与今日谈",
    "cat2": "02_求是网与半月谈",
    "cat3": "03_纽约时报观点",
}

# 栏目定义。list_kind 决定用哪个列表抓取函数。
SOURCES = {
    "rmsp": {
        "label": "人民时评",
        "category": "cat1",
        "list_kind": "people",
        "list_url": "http://opinion.people.com.cn/GB/8213/49160/49219/index.html",
    },
    "jrt": {
        "label": "今日谈",
        "category": "cat1",
        "list_kind": "people",
        "list_url": "http://opinion.people.com.cn/GB/8213/49160/49221/index.html",
    },
    "qs": {
        "label": "求是·社论评论",
        "category": "cat2",
        "list_kind": "qiushi",
        "list_url": "https://www.qstheory.cn/v9zhuanqu/zhuanqu/slpl/index.htm",
    },
    "byt": {
        "label": "半月谈·评论",
        "category": "cat2",
        "list_kind": "banyuetan",
        "list_url": "http://www.banyuetan.org/byt/banyuetanpinglun/index.html",
    },
    "nyt": {
        "label": "纽约时报·观点",
        "category": "cat3",
        "list_kind": "nyt",
        "list_url": DEFAULT_NYT_RSS,
    },
}

DEFAULT_SOURCES = ["rmsp", "jrt", "qs", "byt", "nyt"]


# --------------------------------------------------------------------------
# 数据模型
# --------------------------------------------------------------------------
@dataclass
class Item:
    """列表页上的一条候选文章。"""

    url: str
    title: str
    date: str = ""  # YYYY-MM-DD，尽力解析，解析不了则为空


@dataclass
class Article:
    """抓取完成、待导出的文章。"""

    source_id: str
    source_label: str
    category: str
    url: str
    title: str
    date: str = ""
    author: str = ""
    source_text: str = ""
    summary: str = ""
    paragraphs: list = field(default_factory=list)
    notes: list = field(default_factory=list)


# --------------------------------------------------------------------------
# 通用工具
# --------------------------------------------------------------------------
def log(msg: str) -> None:
    print(msg, flush=True)


def clean_text(text: str) -> str:
    """合并空白字符，去掉首尾空格。"""
    if not text:
        return ""
    return re.sub(r"\s+", " ", text).strip(" \u3000")


def normalize_date(raw: str) -> str:
    """把常见日期写法转成 YYYY-MM-DD，转不了返回空字符串。"""
    if not raw:
        return ""
    raw = raw.strip()
    m = re.search(r"(\d{4})[年/\-.](\d{1,2})[月/\-.](\d{1,2})", raw)
    if m:
        y, mo, d = (int(x) for x in m.groups())
        try:
            return dt.date(y, mo, d).isoformat()
        except ValueError:
            return ""
    m = re.search(r"\b(\d{4})(\d{2})(\d{2})\b", raw)
    if m:
        y, mo, d = (int(x) for x in m.groups())
        try:
            return dt.date(y, mo, d).isoformat()
        except ValueError:
            return ""
    return ""


def strip_column_tag(title: str, label: str) -> str:
    """去掉标题末尾的栏目括号，如 “（人民时评）”。"""
    pat = re.compile(r"[（(]\s*" + re.escape(label) + r"\s*[)）]\s*$")
    return pat.sub("", clean_text(title))


def safe_filename(name: str, max_len: int = 90) -> str:
    name = clean_text(name)
    name = re.sub(r'[\\/:*?"<>|\r\n\t]', "_", name)
    name = re.sub(r"\s+", " ", name).strip(" .")
    if len(name) > max_len:
        name = name[: max_len - 1].rstrip() + "…"
    return name or "未命名文章"


class FetchError(Exception):
    pass


def _request(
    session: requests.Session,
    url: str,
    timeout: int = 30,
    attempts: int = 3,
) -> requests.Response:
    last_err: Exception | None = None
    for attempt in range(attempts):
        try:
            resp = session.get(url, timeout=timeout)
            resp.raise_for_status()
            return resp
        except Exception as exc:  # noqa: BLE001
            last_err = exc
            time.sleep(1.5 * (attempt + 1))
    raise FetchError(f"请求失败：{url}\n  原因：{last_err}") from last_err


def get_soup(
    session: requests.Session,
    url: str,
    timeout: int = 30,
    attempts: int = 3,
) -> BeautifulSoup:
    resp = _request(session, url, timeout=timeout, attempts=attempts)
    if not resp.encoding or resp.encoding.lower() in ("iso-8859-1", "us-ascii"):
        resp.encoding = resp.apparent_encoding or "utf-8"
    return BeautifulSoup(resp.text, "html.parser")


def filter_item_date(item: Item, cutoff: dt.date | None) -> bool:
    if cutoff is None or not item.date:
        return True
    try:
        return dt.date.fromisoformat(item.date) >= cutoff
    except ValueError:
        return True


# --------------------------------------------------------------------------
# 列表页抓取
# --------------------------------------------------------------------------
def fetch_people_items(
    session: requests.Session,
    source_cfg: dict,
    want: int,
    cutoff: dt.date | None,
    max_pages: int,
) -> list[Item]:
    """人民日报言论库栏目列表：30 条/页，可翻页，时间倒序。"""
    items: list[Item] = []
    seen: set[str] = set()
    page_url = source_cfg["list_url"]
    page_no = 0
    while page_url and len(items) < want and page_no < max_pages:
        page_no += 1
        try:
            soup = get_soup(session, page_url)
        except FetchError:
            break
        lis = soup.select("div.leftItem ul li")
        for li in lis:
            a = li.find("a", href=True)
            if not a:
                continue
            href = urljoin(page_url, a["href"])
            if not re.search(r"/n1/20\d{2}/\d{4}/c\d+-\d+\.html", href):
                continue
            if href in seen:
                continue
            date_el = li.find("i")
            item = Item(
                url=href,
                title=clean_text(a.get_text(" ", strip=True)),
                date=normalize_date(date_el.get_text(" ", strip=True) if date_el else ""),
            )
            if filter_item_date(item, cutoff):
                items.append(item)
            seen.add(href)
        next_a = None
        for a in soup.find_all("a", href=True):
            txt = "".join(a.stripped_strings)
            if txt == "下一页":
                next_a = a
                break
        page_url = urljoin(page_url, next_a["href"]) if next_a else None
    return items


def fetch_qiushi_items(
    session: requests.Session,
    source_cfg: dict,
    want: int,
    cutoff: dt.date | None,
    max_pages: int,
) -> list[Item]:
    """求是网社论评论：单页内联约 50 条，无分页解析。"""
    soup = get_soup(session, source_cfg["list_url"])
    items: list[Item] = []
    seen: set[str] = set()
    for li in soup.select("ul.wz-list li"):
        a = li.find("a", href=True)
        if not a:
            continue
        href = urljoin(source_cfg["list_url"], a["href"])
        if href in seen:
            continue
        seen.add(href)
        spans = li.select(".list-style1-info span")
        mmdd_raw = clean_text(spans[-1].get_text(" ", strip=True)) if spans else ""
        mmdd = normalize_date(mmdd_raw)
        year_m = re.search(r"/(20\d{2})(\d{2})(\d{2})/", href)
        date = ""
        if re.fullmatch(r"\d{1,2}-\d{1,2}", mmdd_raw) and year_m:
            date = f"{year_m.group(1)}-{mmdd_raw}"
        elif mmdd and year_m:
            date = mmdd
        elif year_m:
            date = f"{year_m.group(1)}-{year_m.group(2)}-{year_m.group(3)}"
        item = Item(url=href, title=clean_text(a.get_text("", strip=True)), date=date)
        if filter_item_date(item, cutoff):
            items.append(item)
    return items[:want]


def _banyuetan_page(base_url: str, page_no: int) -> str:
    if page_no <= 1:
        return base_url
    if base_url.endswith("index.html"):
        return base_url[: -len("index.html")] + f"index_{page_no}.html"
    return base_url


def fetch_banyuetan_items(
    session: requests.Session,
    source_cfg: dict,
    want: int,
    cutoff: dt.date | None,
    max_pages: int,
) -> list[Item]:
    """半月谈评论：li > h3 > a，可翻页。"""
    items: list[Item] = []
    seen: set[str] = set()
    for page_no in range(1, max_pages + 1):
        if len(items) >= want:
            break
        url = _banyuetan_page(source_cfg["list_url"], page_no)
        try:
            soup = get_soup(session, url)
        except FetchError:
            break
        found_any = False
        for h3 in soup.find_all("h3"):
            a = h3.find("a", href=True)
            if not a or "/detail/" not in a["href"]:
                continue
            found_any = True
            href = urljoin(url, a["href"])
            if href in seen:
                continue
            seen.add(href)
            li = h3.find_parent("li")
            date_el = li.select_one("span.tag3") if li else None
            item = Item(
                url=href,
                title=clean_text(a.get_text(" ", strip=True)),
                date=normalize_date(date_el.get_text(" ", strip=True) if date_el else ""),
            )
            if filter_item_date(item, cutoff):
                items.append(item)
        if not found_any:
            break
    return items


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _xml_child(el: ET.Element, *names: str) -> ET.Element | None:
    for child in el:
        if _local_name(child.tag) in names:
            return child
    return None


def _xml_text(el: ET.Element | None) -> str:
    return clean_text(el.text or "") if el is not None else ""


def _load_rss_content(session: requests.Session, rss_url: str) -> str:
    """读取 RSS：支持 http(s) 地址，也支持本地 XML 文件路径。"""
    if re.match(r"^https?://", rss_url, re.I):
        return _request(session, rss_url).text
    return Path(rss_url).read_text(encoding="utf-8", errors="replace")


def _parse_nyt_items(content: str, want: int, cutoff: dt.date | None) -> list[Item]:
    root = ET.fromstring(content)
    entries = [e for e in root.iter() if _local_name(e.tag) in ("item", "entry")]
    items: list[Item] = []
    seen: set[str] = set()
    for entry in entries:
        title = _xml_text(_xml_child(entry, "title"))
        link_el = _xml_child(entry, "link")
        link = ""
        if link_el is not None:
            link = link_el.get("href") or _xml_text(link_el)
        pub = _xml_child(entry, "pubDate", "published", "updated")
        date = ""
        if pub is not None:
            raw = clean_text(pub.text or "")
            try:
                date = parsedate_to_datetime(raw).astimezone().date().isoformat()
            except (TypeError, ValueError):
                date = normalize_date(raw)
        item = Item(url=clean_text(link), title=title, date=date)
        if item.url and item.url not in seen and filter_item_date(item, cutoff):
            items.append(item)
            seen.add(item.url)
        if len(items) >= want:
            break
    return items


def _parse_nyt_summaries(content: str) -> dict[str, str]:
    """从 RSS 内容中提取 description，作为全文失败时的摘要回退。"""
    try:
        root = ET.fromstring(content)
    except ET.ParseError:
        return {}
    out: dict[str, str] = {}
    for entry in root.iter():
        if _local_name(entry.tag) not in ("item", "entry"):
            continue
        link_el = _xml_child(entry, "link")
        link = ""
        if link_el is not None:
            link = link_el.get("href") or _xml_text(link_el)
        desc_el = _xml_child(entry, "description", "summary", "content")
        desc = ""
        if desc_el is not None and desc_el.text:
            desc = clean_text(
                BeautifulSoup(desc_el.text, "html.parser").get_text(" ", strip=True)
            )
        if link and desc:
            out[clean_text(link)] = desc
    return out


def fetch_nyt_items(
    session: requests.Session,
    source_cfg: dict,
    want: int,
    cutoff: dt.date | None,
    max_pages: int,
    rss_override: str = "",
    rss_content: str | None = None,
) -> list[Item]:
    """NYT Opinion RSS。rss_override 可以是 http(s) 地址或本地 XML 文件路径。"""
    rss_url = rss_override or source_cfg["list_url"]
    if rss_content is None:
        rss_content = _load_rss_content(session, rss_url)
    return _parse_nyt_items(rss_content, want, cutoff)


LIST_FETCHERS = {
    "people": fetch_people_items,
    "qiushi": fetch_qiushi_items,
    "banyuetan": fetch_banyuetan_items,
    "nyt": fetch_nyt_items,
}


# --------------------------------------------------------------------------
# 正文抓取
# --------------------------------------------------------------------------
def _body_paragraphs(container: BeautifulSoup | None) -> list[str]:
    if container is None:
        return []
    paras: list[str] = []
    junk_exact = {
        "分享让更多人看到",
        "点击播报本文，约",
        "纠错",
        "字号",
        "打印本页",
        "关闭窗口",
        "相关阅读",
    }
    for p in container.find_all("p"):
        text = clean_text(p.get_text(" ", strip=True))
        if not text or len(text) < 3:
            continue
        if text in junk_exact or text.startswith("点击播报本文"):
            continue
        if re.match(r"^(责任编辑|责编|编辑|执笔|版权|原标题|点击下载|免责声明)[:：]", text):
            continue
        if re.search(r"《\s*人民日报\s*》\s*（?\s*\d{4}年\d{1,2}月\d{1,2}日\s*\d{1,2}\s*版", text):
            continue
        paras.append(text)
    return paras


def _pick_article_h1(soup: BeautifulSoup) -> str:
    for h1 in soup.find_all("h1"):
        if h1.find("img"):
            continue
        text = clean_text(h1.get_text(" ", strip=True))
        if len(text) >= 4:
            return text
    return ""


def _fetch_people_article(
    session: requests.Session, source_cfg: dict, item: Item, source_id: str
) -> Article:
    soup = get_soup(session, item.url)
    title = _pick_article_h1(soup) or item.title
    date_el = soup.find(id="newstime")
    date = normalize_date(date_el.get_text(" ", strip=True) if date_el else "") or item.date
    meta_source = soup.find("meta", attrs={"name": "source"})
    source_text = clean_text(meta_source.get("content", "")) if meta_source else ""
    if source_text.startswith("来源："):
        source_text = source_text[len("来源：") :]
    body = _body_paragraphs(soup.select_one("div.rm_txt_con"))
    if not body:
        raise FetchError(f"未解析到正文：{item.url}")
    return Article(
        source_id=source_id,
        source_label=source_cfg["label"],
        category=source_cfg["category"],
        url=item.url,
        title=strip_column_tag(title, source_cfg["label"]),
        date=date,
        source_text=source_text or "人民日报",
        paragraphs=body,
    )


def _fetch_qiushi_article(
    session: requests.Session, source_cfg: dict, item: Item
) -> Article:
    soup = get_soup(session, item.url)
    title = _pick_article_h1(soup) or item.title
    meta_date = soup.find("meta", attrs={"name": "publishdate"})
    date = normalize_date(meta_date.get("content", "") if meta_date else "") or item.date
    author = ""
    source_text = ""
    h2 = soup.select_one("div.content h2") or soup.select_one("h2")
    if h2:
        for span in h2.find_all("span"):
            txt = clean_text(span.get_text(" ", strip=True))
            if txt.startswith("作者-"):
                author = txt[len("作者-") :]
            elif txt.startswith("来源-"):
                source_text = txt[len("来源-") :]
    body_el = soup.select_one("#detailContent") or soup.select_one("#detail")
    paras = _body_paragraphs(body_el)
    # 去掉正文开头重复的标题行/署名行
    title_n = clean_text(title)
    author_n = clean_text(author)
    cleaned: list[str] = []
    for p in paras:
        if p == title_n:
            continue
        if author_n and p == author_n:
            continue
        if not author_n and re.match(r"^《求是》.*评论员$", p):
            author = p
            continue
        cleaned.append(p)
    if not cleaned:
        raise FetchError(f"未解析到正文：{item.url}")
    return Article(
        source_id="qs",
        source_label=source_cfg["label"],
        category=source_cfg["category"],
        url=item.url,
        title=title,
        date=date,
        author=author,
        source_text=source_text or "求是网",
        paragraphs=cleaned,
    )


def _fetch_banyuetan_article(
    session: requests.Session, source_cfg: dict, item: Item
) -> Article:
    soup = get_soup(session, item.url)
    h1 = soup.select_one("div.detail_tit h1")
    title = clean_text(h1.get_text(" ", strip=True)) if h1 else _pick_article_h1(soup)
    title = title or item.title
    time_el = soup.select_one("div.detail_tit_time")
    date = normalize_date(time_el.get_text(" ", strip=True) if time_el else "") or item.date
    src_el = soup.select_one("div.detail_tit_source")
    source_text = clean_text(src_el.get_text(" ", strip=True)) if src_el else ""
    if source_text.startswith("来源："):
        source_text = source_text[len("来源：") :]
    meta_author = soup.find("meta", attrs={"name": "author"})
    author = clean_text(meta_author.get("content", "")) if meta_author else ""
    paras = _body_paragraphs(soup.select_one("div.detail_content"))
    cleaned: list[str] = []
    for p in paras:
        m = re.match(r"^(半月谈(评论员|记者|观察员)?|文[/·]?\s*)[:：]?\s*(.+)$", p)
        if (
            m
            and len(p) <= 30
            and not re.search(r"[。；，、！？]", p)
        ):
            author = m.group(3)
            continue
        if re.match(r"^编辑[:：]", p):
            continue
        cleaned.append(p)
    if not cleaned:
        raise FetchError(f"未解析到正文：{item.url}")
    label = "半月谈·今日谈" if "/jrt/detail/" in item.url else source_cfg["label"]
    return Article(
        source_id="byt",
        source_label=label,
        category=source_cfg["category"],
        url=item.url,
        title=title,
        date=date,
        author=author,
        source_text=source_text or "半月谈网",
        paragraphs=cleaned,
    )


def _parse_nyt_page(soup: BeautifulSoup) -> tuple[str, str, str, list[str]]:
    """返回 (标题, 作者, 摘要, 正文段落)。无法取得正文时抛出 FetchError。"""
    title = ""
    og = soup.find("meta", attrs={"property": "og:title"})
    if og:
        title = clean_text(og.get("content", ""))
    title = title or _pick_article_h1(soup)
    byline = ""
    meta_by = soup.find("meta", attrs={"name": "byl"})
    if meta_by:
        byline = clean_text(meta_by.get("content", ""))
    if not byline:
        bl = soup.select_one('div[data-testid="byline"]') or soup.select_one(
            'p[class*="byline"]'
        )
        if bl:
            byline = clean_text(bl.get_text(" ", strip=True))
    desc = ""
    meta_desc = soup.find("meta", attrs={"property": "og:description"}) or soup.find(
        "meta", attrs={"name": "description"}
    )
    if meta_desc:
        desc = clean_text(meta_desc.get("content", ""))
    selectors = [
        'div[data-testid="article-body"]',
        'section[name="articleBody"]',
        'div[data-name="articleBody"]',
        'div.StoryBodyCompanionColumn',
    ]
    paras: list[str] = []
    for sel in selectors:
        node = soup.select_one(sel)
        if not node:
            continue
        paras = [
            clean_text(p.get_text(" ", strip=True))
            for p in node.find_all("p")
            if clean_text(p.get_text(" ", strip=True))
        ]
        if len("".join(paras)) > 300:
            break
        paras = []
    if len("".join(paras)) < 300:
        raise FetchError("完整正文不可见（可能受订阅/访问限制）")
    return title, byline, desc, paras


def _fetch_nyt_article(
    session: requests.Session, source_cfg: dict, item: Item, summary: str = ""
) -> Article:
    try:
        soup = get_soup(session, item.url, timeout=15, attempts=1)
        title, byline, desc, paras = _parse_nyt_page(soup)
        notes: list[str] = []
        article = Article(
            source_id="nyt",
            source_label=source_cfg["label"],
            category=source_cfg["category"],
            url=item.url,
            title=title or item.title,
            date=item.date,
            author=byline,
            source_text="The New York Times",
            summary=desc,
            paragraphs=paras,
            notes=notes,
        )
        return article
    except FetchError as exc:
        # 拿不到全文时退化为 RSS 摘要，保证文件能生成。
        note = re.sub(r"\s+", " ", f"未能取得全文（{exc}），本文件仅含 RSS 摘要，可通过链接在线阅读。")
        return Article(
            source_id="nyt",
            source_label=source_cfg["label"],
            category=source_cfg["category"],
            url=item.url,
            title=item.title,
            date=item.date,
            source_text="The New York Times",
            summary=summary,
            notes=[note],
        )


def download_article(
    session: requests.Session,
    source_id: str,
    source_cfg: dict,
    item: Item,
    rss_summaries: dict[str, str],
) -> Article:
    if source_id == "rmsp" or source_id == "jrt":
        return _fetch_people_article(session, source_cfg, item, source_id)
    if source_id == "qs":
        return _fetch_qiushi_article(session, source_cfg, item)
    if source_id == "byt":
        return _fetch_banyuetan_article(session, source_cfg, item)
    if source_id == "nyt":
        return _fetch_nyt_article(session, source_cfg, item, rss_summaries.get(item.url, ""))
    raise FetchError(f"未知来源：{source_id}")


# --------------------------------------------------------------------------
# Word 导出
# --------------------------------------------------------------------------
def _docx_font(run, east: str = "宋体", latin: str = "Times New Roman") -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)


def export_docx(article: Article, path: Path) -> None:
    if not HAVE_DOCX:
        raise FetchError("缺少 python-docx，请先执行 pip install -r requirements.txt")
    document = docx.Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_para(
        text: str,
        size: float = 12,
        bold: bool = False,
        gray: bool = False,
        indent: bool = False,
        align: str = "left",
    ) -> None:
        p = document.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5
        if indent:
            pf.first_line_indent = Pt(size * 2)
        if align == "center":
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "justify":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        _docx_font(run)
        run.font.size = Pt(size)
        run.bold = bold
        if gray:
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_para(article.title or "（无标题）", size=20, bold=True, align="center")
    add_para("", size=6)
    meta = []
    meta.append(f"栏目：{article.source_label}")
    if article.source_text:
        meta.append(f"来源：{article.source_text}")
    if article.author:
        meta.append(f"作者：{article.author}")
    if article.date:
        meta.append(f"日期：{article.date}")
    add_para("　".join(meta), size=10.5, gray=True)
    add_para(f"原文链接：{article.url}", size=10.5, gray=True)
    add_para("", size=6)
    for note in article.notes:
        add_para(f"【提示】{note}", size=10.5, gray=True)
    if article.paragraphs:
        for para in article.paragraphs:
            add_para(para, size=12, align="justify", indent=True)
    elif article.summary:
        add_para(article.summary, size=12, align="justify", indent=True)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


# --------------------------------------------------------------------------
# PDF 导出（PyMuPDF 内置简体中文字体 china-s）
# --------------------------------------------------------------------------
PDF_FONT = "china-s"
PDF_PAGE_W = 595.0
PDF_PAGE_H = 842.0
PDF_MARGIN = 50.0


def export_pdf(article: Article, path: Path) -> None:
    if fitz is None:
        raise FetchError("缺少 PyMuPDF，请先执行 pip install -r requirements.txt")
    doc = fitz.open()
    page = doc.new_page(width=PDF_PAGE_W, height=PDF_PAGE_H)
    y = PDF_MARGIN
    max_w = PDF_PAGE_W - 2 * PDF_MARGIN

    def text_block(
        text: str,
        size: float = 11.5,
        gap: float = 1.65,
        color: tuple = (0.1, 0.1, 0.1),
        space_after: float = 5.0,
        align: int = 0,
        _depth: int = 0,
    ) -> None:
        nonlocal page, y
        if not text:
            y += space_after
            return
        if hasattr(fitz, "get_text_length"):
            total_w = fitz.get_text_length(text, fontname=PDF_FONT, fontsize=size)
        else:  # pragma: no cover - 兼容性回退
            total_w = len(text) * size
        lines = max(1, math.ceil(total_w / max_w) + 1)
        line_h = size * gap
        need = lines * line_h
        rc = None
        for boost in (1.0, 1.5, 2.4):
            box_h = need * boost + 4
            if y + box_h > PDF_PAGE_H - PDF_MARGIN:
                page = doc.new_page(width=PDF_PAGE_W, height=PDF_PAGE_H)
                y = PDF_MARGIN
            rc = page.insert_textbox(
                fitz.Rect(PDF_MARGIN, y, PDF_MARGIN + max_w, y + box_h),
                text,
                fontname=PDF_FONT,
                fontsize=size,
                lineheight=gap,
                color=color,
                align=align,
            )
            if rc is None or rc >= -0.6:
                break
        if rc is not None and rc < -0.6 and _depth < 3:
            # 极罕见的排版溢出：把段落拆成两半分别排版，避免丢字
            mid = len(text) // 2
            text_block(
                text[:mid],
                size=size,
                gap=gap,
                color=color,
                space_after=0,
                align=align,
                _depth=_depth + 1,
            )
            text_block(
                text[mid:],
                size=size,
                gap=gap,
                color=color,
                space_after=space_after,
                align=align,
                _depth=_depth + 1,
            )
            return
        y += (need if rc is None or rc >= -0.6 else need * 2.4) + space_after

    text_block(article.title or "（无标题）", size=17, gap=1.35, space_after=10)
    meta = f"栏目：{article.source_label}"
    if article.source_text:
        meta += f"　来源：{article.source_text}"
    if article.author:
        meta += f"　作者：{article.author}"
    if article.date:
        meta += f"　日期：{article.date}"
    text_block(meta, size=10, gap=1.4, color=(0.45, 0.45, 0.45), space_after=3)
    text_block(
        f"原文链接：{article.url}",
        size=9.5,
        gap=1.4,
        color=(0.35, 0.45, 0.65),
        space_after=10,
    )
    for note in article.notes:
        text_block(f"【提示】{note}", size=10, gap=1.4, color=(0.7, 0.35, 0.1), space_after=8)
    if article.paragraphs:
        for para in article.paragraphs:
            text_block(para, size=11.5, space_after=6)
    elif article.summary:
        text_block(article.summary, size=11.5, space_after=6)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


# --------------------------------------------------------------------------
# 主流程
# --------------------------------------------------------------------------
def _build_session() -> requests.Session:
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": USER_AGENT,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.7",
        }
    )
    return session


def _gather_items(
    session: requests.Session,
    source_id: str,
    source_cfg: dict,
    want: int,
    cutoff: dt.date | None,
    max_pages: int,
    nyt_rss: str,
    nyt_rss_content: str | None = None,
) -> list[Item]:
    kind = source_cfg["list_kind"]
    fetcher = LIST_FETCHERS[kind]
    if kind == "nyt":
        if nyt_rss_content is None:
            raise FetchError("NYT RSS 读取失败或为空")
        items = fetcher(
            session,
            source_cfg,
            want,
            cutoff,
            max_pages,
            rss_override=nyt_rss,
            rss_content=nyt_rss_content,
        )
    else:
        items = fetcher(session, source_cfg, want, cutoff, max_pages)
    return items[:want]


def run(args: argparse.Namespace) -> int:
    global last_run_summary
    if args.list_sources:
        for sid in DEFAULT_SOURCES:
            cfg = SOURCES[sid]
            log(f"{sid:<6} {cfg['label']:<12} -> {CATEGORIES[cfg['category']]}")
        return 0

    source_ids = [s.strip() for s in re.split(r"[,\s]+", args.sources) if s.strip()]
    unknown = [s for s in source_ids if s not in SOURCES]
    if unknown:
        log(f"未知栏目：{', '.join(unknown)}；可用栏目：{', '.join(SOURCES)}")
        return 2
    if not source_ids:
        source_ids = list(DEFAULT_SOURCES)

    fmt = args.format
    if fmt in ("pdf", "both") and fitz is None:
        log("错误：未安装 PyMuPDF，无法导出 PDF。请执行：pip install -r requirements.txt")
        return 2
    if fmt in ("docx", "both") and not HAVE_DOCX:
        log("错误：未安装 python-docx，无法导出 Word。请执行：pip install -r requirements.txt")
        return 2

    out_root = Path(args.out).expanduser()
    out_root.mkdir(parents=True, exist_ok=True)
    for cat_dir in CATEGORIES.values():
        (out_root / cat_dir).mkdir(parents=True, exist_ok=True)
    cutoff = None
    if args.days > 0:
        cutoff = dt.date.today() - dt.timedelta(days=args.days)

    session = _build_session()
    nyt_rss_content: str | None = None
    nyt_summaries: dict[str, str] = {}
    if "nyt" in source_ids:
        try:
            nyt_rss_content = _load_rss_content(session, args.nyt_rss)
            nyt_summaries = _parse_nyt_summaries(nyt_rss_content)
        except Exception as exc:  # noqa: BLE001
            log(f"提示：NYT RSS 读取失败（{exc}），NYT 栏目将被跳过")

    log("=" * 62)
    log("时评文章下载归档工具")
    log(f"输出目录：{out_root}")
    log(f"栏目：{', '.join(source_ids)}　每栏目 {args.count} 篇　格式：{fmt}")
    log("=" * 62)

    errors: list[str] = []
    written = 0
    for idx, source_id in enumerate(source_ids, 1):
        cfg = SOURCES[source_id]
        label = cfg["label"]
        try:
            items = _gather_items(
                session,
                source_id,
                cfg,
                args.count,
                cutoff,
                args.max_pages,
                args.nyt_rss,
                nyt_rss_content,
            )
        except Exception as exc:  # noqa: BLE001
            errors.append(f"[{label}] 获取文章列表失败：{exc}")
            log(f"[{idx}/{len(source_ids)}] ✗ {label}：列表获取失败\n    {exc}")
            continue
        log(f"[{idx}/{len(source_ids)}] {label}：共获取 {len(items)} 篇，开始下载正文…")
        for j, item in enumerate(items, 1):
            try:
                article = download_article(session, source_id, cfg, item, nyt_summaries)
            except Exception as exc:  # noqa: BLE001
                errors.append(f"[{label}] {item.title} 正文失败：{exc}")
                log(f"    - 第 {j} 篇《{item.title}》下载失败：{exc}")
                continue
            cat_dir = out_root / CATEGORIES[article.category]
            date_part = re.sub(r"-", "", article.date) if article.date else "no-date"
            base = f"{date_part}_{article.source_label}_{safe_filename(article.title)}"
            files: list[Path] = []
            for ext in ([fmt] if fmt != "both" else ["pdf", "docx"]):
                dest = cat_dir / f"{base}.{ext}"
                try:
                    if ext == "pdf":
                        export_pdf(article, dest)
                    else:
                        export_docx(article, dest)
                    files.append(dest)
                except Exception as exc:  # noqa: BLE001
                    errors.append(f"[{label}] {item.title} 导出 .{ext} 失败：{exc}")
                    log(f"    - 导出失败：{exc}")
            written += len(files)
            log(f"    ✓ {article.date or '日期未知'}《{article.title}》 -> "
                f"{', '.join(str(f) for f in files)}")
            time.sleep(max(0.0, args.delay))

    log("-" * 62)
    log(f"完成：共生成 {written} 个文件")
    last_run_summary = f"完成：共生成 {written} 个文件"
    if errors:
        log(f"共 {len(errors)} 条失败记录：")
        for err in errors:
            log("  - " + str(err).replace("\n", " ")[:300])
        last_run_summary += f"，另有 {len(errors)} 条失败记录"
        return 1
    return 0


def build_parser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(
        description="下载人民日报/求是/半月谈/NYT 观点文章并按类别归档到桌面",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    ap.add_argument(
        "-f",
        "--format",
        choices=["pdf", "docx", "both"],
        default="both",
        help="导出格式：pdf / docx / 两者都要",
    )
    ap.add_argument(
        "-n",
        "--count",
        type=int,
        default=5,
        help="每个栏目下载最新的多少篇",
    )
    ap.add_argument(
        "-d",
        "--days",
        type=int,
        default=0,
        help="只下载最近 N 天的文章；0 表示不限",
    )
    ap.add_argument(
        "-s",
        "--sources",
        default=",".join(DEFAULT_SOURCES),
        help="栏目，逗号分隔：" + ", ".join(SOURCES),
    )
    ap.add_argument(
        "-o",
        "--out",
        default=str(DEFAULT_OUT),
        help="输出根目录（默认桌面“时评文章下载”）",
    )
    ap.add_argument(
        "--nyt-rss",
        default=DEFAULT_NYT_RSS,
        help="NYT Opinion RSS 地址（也可以是本地 XML 文件路径，便于离线测试）",
    )
    ap.add_argument(
        "--delay",
        type=float,
        default=0.8,
        help="抓取每篇文章之间的间隔秒数，请勿设太小",
    )
    ap.add_argument(
        "--max-pages",
        type=int,
        default=6,
        help="每个栏目最多翻页数",
    )
    ap.add_argument("--list-sources", action="store_true", help="仅列出可用栏目")
    return ap


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        return run(args)
    except KeyboardInterrupt:
        log("\n已中断。")
        return 130
    except Exception as exc:  # noqa: BLE001
        log(f"运行出错：{exc}")
        if os.environ.get("ART_DL_DEBUG"):
            traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
